from google import genai
from google.genai import types
import pandas as pd
from sentence_transformers import SentenceTransformer, util
import fitz
from docx import Document
import pathlib
import os

def extract_docx(filepath):
    
    doc = Document(filepath)

    text = []
    
    #paragraphs
    for para in doc.paragraphs:
        text.append(para.text)
    
    #tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    
    #headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            text.append(para.text)
        for para in section.footer.paragraphs:
            text.append(para.text)


    full_text = "\n".join([t for t in text if t.strip() != ""])
    return full_text

def resume_fitment(resume_filepath, jd_filepath):

    os.environ["TOKENIZERS_PARALLELISM"] = "false"

    #sbert
    model = SentenceTransformer("all-mpnet-base-v2")

    
    if ".pdf" in resume_filepath:
        filepath = pathlib.Path(resume_filepath)
        resume_part = types.Part.from_bytes(data=filepath.read_bytes(), mime_type="application/pdf")

        # used for semantic embeddings comparison
        doc = fitz.open(resume_filepath)
        resume_text = ""
        for page in doc:
            resume_text += page.get_text()

    elif ".docx" in resume_filepath:
        resume_text = extract_docx(resume_filepath)
        resume_part = types.Part.from_text(text=resume_text)

    if ".pdf" in jd_filepath:
        filepath = pathlib.Path(jd_filepath)
        jd_part = types.Part.from_bytes(data=filepath.read_bytes(), mime_type="application/pdf")
        
        # used for semantic embeddings comparison
        doc_jd = fitz.open(jd_filepath)
        jd_text = ""
        for page in doc_jd:
            jd_text += page.get_text()
    
    elif ".docx" in jd_filepath:
        jd_text = extract_docx(jd_filepath)
        jd_part = types.Part.from_text(text=jd_text)



    #get job description text
    #df = pd.read_csv("/Users/aahan_bagga/Downloads/job_descriptions.csv", encoding="latin1")

    # convert to vectors

    short_vec = model.encode(resume_text, convert_to_tensor = True)
    long_vec = model.encode(jd_text, convert_to_tensor = True)

    #cosine similarity will tell us the angle between those two vectors in the high dimensional vector space; that is our rating

    score = util.cos_sim(short_vec, long_vec).item()

    prompt = "Print the rating for this candidate from 1 to 5 stars, the similarity score, and whether the hire should be considered or not in the most descriptive terms possible; if there is a skill gap worth investment, note it in the recommendation. Then the justification underneath with exactly six sentences. If there's a skill gap worth employer investment, mention it." \
            "For example, if the candidate meets ~75% of required skills but shows other valuable traits, note this. Only mention the skill gap if" \
            "it’s the sole concern. Do not skip any reasoning. Use the provided SBERT similarity score in your logic as well: "
    
    #expecting job description as raw text, can change as needed
    #job_description = "We are seeking a Frontend Engineer with a strong command of Flutter and Dart for building responsive, scalable cross-platform applications." \
     #               "The ideal candidate has at least 3 years of experience working with Flutter and is comfortable implementing both frontend UI components and backend logic within a unified Dart codebase." \
      #              "You will be responsible for crafting rich mobile and web experiences that consume data from RESTful APIs, handle asynchronous HTTP requests," \
       #             "and manage application state effectively using tools such as Provider, Riverpod, or BLoC.This role requires a solid understanding of client-server communication, JSON serialization," \
        #            "and best practices in data modeling, caching, and error handling. Familiarity with backend-in-Flutter patterns (such as cloud functions, local server integrations, or Firebase)" \
         #           "is a strong plus. You’ll collaborate closely with product managers and designers to deliver intuitive, performant apps that work seamlessly across devices."

    
    #LLM gemini 2.0 flash
    client = genai.Client(api_key ="AIzaSyC_hptz6aVNaanku8pfAArrpaUdN5VkfCk")
    response = client.models.generate_content(
    model="gemini-2.0-flash",
    contents=[resume_part, prompt+str(score), jd_part])

    return response.text

print(resume_fitment('/Users/aahan_bagga/Downloads/642135000000401274_Profile Data Engineer_Prashant Srivastava.pdf', '/Users/aahan_bagga/Downloads/Data_scientist.pdf'))