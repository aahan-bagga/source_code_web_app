from flask import Flask, request, jsonify
from flask_cors import CORS
from openai import OpenAI
from sentence_transformers import SentenceTransformer, util
from docx import Document
import fitz
import os
import pathlib
import json
import re
import pandas as pd
import pypandoc

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024
CORS(app)

os.environ["TOKENIZERS_PARALLELISM"] = "false"

# Load SBERT model and OpenAI GPT client
model = SentenceTransformer("all-mpnet-base-v2")
openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
GPT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-5-nano-2025-08-07")

def extract_docx(filepath): #separate DOCX parser
    doc = Document(filepath)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    for section in doc.sections:
        for para in section.header.paragraphs:
            text.append(para.text)
        for para in section.footer.paragraphs:
            text.append(para.text)
    return "\n".join([t for t in text if t.strip() != ""])

def extract_text_and_part(filepath):
    ext = pathlib.Path(filepath).suffix.lower()
    text = ""
    
    if ext == ".pdf":
        doc = fitz.open(filepath)
        text = "\n".join(page.get_text() for page in doc)
    
    elif ext == ".docx":
        text = extract_docx(filepath)
    
    elif ext == ".doc":
        try:
            # Converts .doc to plain text using pandoc
            text = pypandoc.convert_file(filepath, 'plain', format='doc')
        except Exception as e:
            # Fallback: sometimes .doc files are just renamed .rtf or .txt
            try:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except:
                raise ValueError(f"Could not read .doc file: {e}")
    
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
    return text, None # Returning None for 'part' as per your original structure
    
@app.route('/score_resumes_ranked', methods=['POST'])
def score_resumes_ranked():
    if 'resumes' not in request.files or 'job' not in request.files:
        return jsonify({"error": "Missing files"}), 400

    job_file = request.files['job']
    job_path = f"/tmp/{job_file.filename}"
    job_file.save(job_path)
    jd_text, jd_part = extract_text_and_part(job_path)  # jd_part unused, kept for structural parity

    resume_files = request.files.getlist('resumes')
    results = []

    for resume_file in resume_files:
        resume_path = f"/tmp/{resume_file.filename}"
        resume_file.save(resume_path)
        resume_text, _ = extract_text_and_part(resume_path)

        resume_vec = model.encode(resume_text, convert_to_tensor=True)
        jd_vec = model.encode(jd_text, convert_to_tensor=True)
        sbert_score = util.cos_sim(resume_vec, jd_vec).item()

        results.append({
            "resume_name": resume_file.filename,
            "resume_text": resume_text,
            "sbert_score": round(sbert_score, 3)
        })

    # === PROMPT ===
    prompt = """
You are TalentMatchAI, a hiring expert for tech roles in Indian IT services.
You are given 5 resumes and a single job description. Your task is to return a valid JSON object that contains the ranking of these resumes from most to least suitable with the given job description based on the following:

1. *Fitment Score (1–10)*  
   – Based on skills, experience, and qualifications match.

2. *Selection Decision*  
   – “✅ Selected” or “❌ Rejected”  
   – Add 1–2 sentences explaining the decision.

3. *Skill Gap Analysis*  
   – Count required skills: present vs. missing  
   – Rate present skills: Expert / Proficient / Basic  
   – Format as a table:

   | Skill | Required? | Present? | Depth (Exp/Prof/Basic) |

4. *Relevant Experience Summary*  
   – Total years + domain and tools/tech used.

5. *Skill Presence*  
   – List key JD skills: YES if present, NO if not.

6. *Suggested Domains*
   – 2–3 IT service domains (e.g., BFSI, E‑Commerce) with a 1-2 line reasoning right after explaning why it was suggested. Don't skip the explanation.

Each candidate should have their separate JSON object, here's what each should look like with the following structure:
    Ouptut the following for ONLY THE AMOUNT of candidates INPUTTED: Don't provide any NULL values.

{
  "Ranking": [
    {
      "name": string,
      "sbert_score": float,
      "fitment_score": int,
      "selection": boolean,
      "rationale": string,
      "skill_gap_table": [
        { "skill": string, "required": boolean, "present": boolean, "depth": string or null }
      ],
      "experience_summary": string,
      "skill_presence": {
        "Python": boolean,
        "RESTful API": boolean,
        ...
      },
      "suggested_domains": [string, string, string]
      "resume_filename": string
    },
    ...
  ],
  "Summary": string
}

Include a thorough and concise one to two line explanation for why the suggested domain was chosen in the "suggested_domains" part of the JSON object based on the candidate's given details.
Make sure to include these explanations right next to the domains you suggest for each in the "suggested_domains" part of JSON object.

Include the SBERT similarity score for each candidate in your evaluation.
Output in rank order (1 = best fit, N = worst fit) in your formatted JSON object and return ONLY THE JSON OBJECT, no other commentary or explanation.
""".strip()

    # Build a single user message containing JD and resumes (keeps overall flow similar)
    contents_text = [prompt, "\n=== JOB DESCRIPTION ===\n", jd_text, "\n=== RESUMES ===\n"]
    for r in results:
        resume_blob = (
            f"\n--- RESUME START ---\n"
            f"Resume Filename: {r['resume_name']}\n"
            f"SBERT Score: {r['sbert_score']}\n\n"
            f"{r['resume_text']}\n"
            f"--- RESUME END ---"
        )
        contents_text.append(resume_blob)

    user_message = "\n".join(contents_text)

    # Call GPT with JSON response format
    try:
        response = openai_client.chat.completions.create(
            model=GPT_MODEL,
            temperature=1,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": "You are TalentMatchAI. Return ONLY a strict JSON object per the schema. No prose, no code fences."
                },
                {"role": "user", "content": user_message},
            ],
        )

        raw = response.choices[0].message.content.strip()

        try:  # placing the ranking table in the summary section
            data = json.loads(raw)

            # Format table from ranking data
            rows = []
            for idx, candidate in enumerate(data["Ranking"], start=1):
                rows.append({
                    "Rank": idx,
                    "Candidate Name": candidate["name"],
                    "Fitment Score": f"{candidate['fitment_score']} / 10",
                    "Decision": "✅ Selected" if candidate["selection"] else "❌ Rejected",
                    "Notes": candidate["rationale"]
                })

            df = pd.DataFrame(rows)
            html_table = df.to_html(index=False, classes="ranking-table", border=1) #conversion to HTML

            # Inject the HTML table into the summary section
            data["Summary"] = f"""
            <div>
                <h3>Ranked Candidates</h3>
                {html_table}
                <p><strong>Note:</strong> {data['Summary']}</p>
            </div>
            """

            return jsonify(data)

        except Exception as e:
            return jsonify({
                "error": "Invalid JSON from GPT",
                "raw": raw,
                "exception": str(e)
            }), 500

    except Exception as e:
        return jsonify({
            "error": "GPT API error",
            "exception": str(e)
        }), 500

@app.route('/')
def home():
    return 'Resume Fitment Ranking API is live! (OpenAI GPT)'

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000)
