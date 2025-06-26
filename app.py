from flask import Flask, request, jsonify                         # >>> ADDED: Flask for REST API
from flask_cors import CORS                                       # >>> ADDED: Enable CORS for frontend access
from google import genai
from google.genai import types
from sentence_transformers import SentenceTransformer, util
from docx import Document
import fitz
import os
import pathlib

# >>> ADDED: Initialize Flask app
app = Flask(__name__)
CORS(app)  # >>> ADDED: Allow cross-origin requests for frontend apps

os.environ["TOKENIZERS_PARALLELISM"] = "false"

# >>> MOVED: Model and API client initialization outside function so it’s reused
model = SentenceTransformer("all-mpnet-base-v2")
genai_client = genai.Client(api_key="GEMINI_API_KEY")  # >>> CHANGED: use a named client object instead of inline

# >>> MOVED: this helper was unchanged
def extract_docx(filepath):
    doc = Document(filepath)
    text = []

    for para in doc.paragraphs:
        text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)

    for section in doc.sections:
        for para in section.header.paragraphs:
            text.append(para.text)
        for para in section.footer.paragraphs:
            text.append(para.text)

    return "\n".join([t for t in text if t.strip() != ""])

# >>> ADDED: Utility function that combines your PDF and DOCX logic into one
def extract_text_and_part(filepath):
    ext = pathlib.Path(filepath).suffix.lower()
    if ext == ".pdf":
        part = types.Part.from_bytes(data=pathlib.Path(filepath).read_bytes(), mime_type="application/pdf")
        doc = fitz.open(filepath)
        text = "\n".join(page.get_text() for page in doc)
    elif ext == ".docx":
        text = extract_docx(filepath)
        part = types.Part.from_text(text=text)
    else:
        raise ValueError("Unsupported file type")
    return text, part

# >>> ADDED: Flask API endpoint to handle POST request and files
@app.route('/score_resume', methods=['POST'])
def score_resume():
    if 'resume' not in request.files or 'job' not in request.files:
        return jsonify({"error": "Missing files"}), 400

    # >>> CHANGED: Save uploaded files to temporary directory (Render safe)
    resume_file = request.files['resume']
    job_file = request.files['job']
    resume_path = f"/tmp/{resume_file.filename}"
    job_path = f"/tmp/{job_file.filename}"
    resume_file.save(resume_path)
    job_file.save(job_path)

    # >>> CHANGED: Unified file reading for both PDFs and DOCX
    resume_text, resume_part = extract_text_and_part(resume_path)
    jd_text, jd_part = extract_text_and_part(job_path)

    # >>> UNCHANGED: Use SBERT to encode and compare
    resume_vec = model.encode(resume_text, convert_to_tensor=True)
    jd_vec = model.encode(jd_text, convert_to_tensor=True)
    score = util.cos_sim(resume_vec, jd_vec).item()

    # >>> UNCHANGED: Prompt with embedded similarity score
    prompt = (
        "Print the rating for this candidate from 1 to 5 stars, the similarity score, and whether the hire should be considered or not in the most descriptive terms possible; "
        "if there is a skill gap worth investment, note it in the recommendation. Then the justification underneath with exactly six sentences. "
        "Only mention the skill gap if it’s the sole concern. Do not skip any reasoning. Use the provided SBERT similarity score in your logic as well: "
        + str(score)
    )

    prompt2 = ("""You are TalentMatchAI, a hiring expert for tech roles in Indian IT services. Analyze a candidate’s resume and job description to perform:

1. *Fitment Score (1–10)*  
   – Based on skills, experience, and qualifications match.

2. *Selection Decision*  
   – “✅ Selected” or “❌ Rejected”  
   – Add 1–2 sentences explaining the decision.

3. *Skill Gap Analysis*  
   – Count required skills: present vs. missing  
   – Rate present skills: Expert / Proficient / Basic  
   – Format as a table:

   | Skill | Required? | Present? | Depth (Exp/Prof/Basic) |

4. *Relevant Experience Summary*  
   – Total years + domain and tools/tech used.

5. *Skill Presence*  
   – List key JD skills: ✔️ if present, ❌ if not.

6. *Suggested Domains* (if selected)  
   – 2–3 IT service domains (e.g., BFSI, E‑Commerce) with 1-line reasoning each.

Include the provided SBERT similarity score in your evaluation.

### Response Format:
1. Fitment Score:  
2. Selection:  
3. Rationale:  
4. Skill Gap Table:  
| Skill | Required? | Present? | Depth |
5. Experience Summary:  
6. Skill Presence:  
– skill_1: ✔️, skill_2: ❌  
7. Recommended Domains:  
– Domain 1: reason  
– Domain 2: reason
               
               """)

    # >>> CHANGED: Use shared genai client to call Gemini
    response = genai_client.models.generate_content(
        model="gemini-2.0-flash",
        contents=[resume_part, prompt, jd_part]
    )

    # >>> CHANGED: Return proper API response as JSON
    return jsonify({
        "sbert_score": round(score, 3),
        "llm_assessment": response.text
    })

@app.route('/print_nothing', methods=['GET'])
def print_nothing():
    return "this was a get request known by print_nothing"

@app.route('/')
def home():
    return 'Resume Fitment API is live!'

# >>> CHANGED: Set host and port for Render compatibility
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000)
