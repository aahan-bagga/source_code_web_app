from flask import Flask, request, jsonify
from flask_cors import CORS
from openai import OpenAI
from sentence_transformers import SentenceTransformer, util
from docx import Document
import fitz
import os
import pathlib
import json
import pandas as pd
import werkzeug # For secure filename handling

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024
CORS(app)

os.environ["TOKENIZERS_PARALLELISM"] = "false"

# Load SBERT model and OpenAI GPT client
model = SentenceTransformer("all-mpnet-base-v2")
openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

# Use a stable model alias to avoid 404s on specific dated versions
GPT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

def extract_docx(filepath):
    try:
        doc = Document(filepath)
        text = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                text.extend([cell.text for cell in row.cells])
        return "\n".join([t for t in text if t.strip() != ""])
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {str(e)}")

def extract_text_and_part(filepath):
    ext = pathlib.Path(filepath).suffix.lower()
    
    if ext == ".pdf":
        doc = fitz.open(filepath)
        return "\n".join(page.get_text() for page in doc), None
    
    elif ext == ".docx":
        return extract_docx(filepath), None
    
    elif ext == ".doc" or ext == ".txt":
        # Removed pypandoc dependency as it requires a system-level binary install
        # Falling back to standard UTF-8 read which handles .txt and some .doc layouts
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(), None
        except Exception as e:
            raise ValueError(f"Basic text extraction failed: {str(e)}")
    
    else:
        raise ValueError(f"Unsupported file type: {ext}")

@app.route('/score_resumes_ranked', methods=['POST'])
def score_resumes_ranked():
    if 'resumes' not in request.files or 'job' not in request.files:
        return jsonify({"error": "Missing files"}), 400

    try:
        # Save and extract Job Description
        job_file = request.files['job']
        job_filename = werkzeug.utils.secure_filename(job_file.filename)
        job_path = os.path.join("/tmp", job_filename)
        job_file.save(job_path)
        
        jd_text, _ = extract_text_and_part(job_path)
        
        if not jd_text.strip():
            return jsonify({"error": "Job description file is empty or unreadable"}), 400

        resume_files = request.files.getlist('resumes')
        results = []

        for resume_file in resume_files:
            res_filename = werkzeug.utils.secure_filename(resume_file.filename)
            resume_path = os.path.join("/tmp", res_filename)
            resume_file.save(resume_path)
            
            resume_text, _ = extract_text_and_part(resume_path)
            
            # Vector Similarity
            resume_vec = model.encode(resume_text, convert_to_tensor=True)
            jd_vec = model.encode(jd_text, convert_to_tensor=True)
            sbert_score = util.cos_sim(resume_vec, jd_vec).item()

            results.append({
                "resume_name": resume_file.filename,
                "resume_text": resume_text,
                "sbert_score": round(sbert_score, 3)
            })

    except Exception as e:
        # This captures the "Failed to process" errors specifically
        return jsonify({"error": "File processing failed", "details": str(e)}), 500

    # === UPDATED PROMPT: Explicit JSON Schema to prevent parsing errors ===
    prompt = f"""
You are TalentMatchAI, a hiring expert. Rank these {len(results)} resumes against the Job Description.
Return ONLY a valid JSON object. 

JSON SCHEMA:
{{
  "Ranking": [
    {{
      "name": "string",
      "sbert_score": float,
      "fitment_score": int,
      "selection": boolean,
      "rationale": "string",
      "skill_gap_table": [
        {{ "skill": "string", "required": boolean, "present": boolean, "depth": "string" }}
      ],
      "experience_summary": "string",
      "skill_presence": {{ "Python": boolean, "Snowflake": boolean, "Cortex LLM": boolean }},
      "suggested_domains": [
        {{ "domain": "string", "reasoning": "string" }}
      ],
      "resume_filename": "string"
    }}
  ],
  "Summary": "string"
}}

Note: In 'suggested_domains', provide the reasoning INSIDE the object, not as a separate string next to it.
""".strip()

    contents_text = [prompt, "\n=== JOB DESCRIPTION ===\n", jd_text, "\n=== RESUMES ===\n"]
    for r in results:
        contents_text.append(f"\nFilename: {r['resume_name']}\nSBERT: {r['sbert_score']}\n{r['resume_text']}\n---")

    user_message = "\n".join(contents_text)

    try:
        response = openai_client.chat.completions.create(
            model=GPT_MODEL,
            temperature=0.2, # Lower temperature for stricter JSON compliance
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You are a recruitment AI that only outputs strict JSON."},
                {"role": "user", "content": user_message},
            ],
        )

        raw = response.choices[0].message.content.strip()
        data = json.loads(raw)

        # Build HTML table for the 'Summary' section
        rows = []
        for idx, candidate in enumerate(data["Ranking"], start=1):
            rows.append({
                "Rank": idx,
                "Candidate Name": candidate.get("name", "Unknown"),
                "Fitment Score": f"{candidate.get('fitment_score', 0)} / 10",
                "Decision": "✅ Selected" if candidate.get("selection") else "❌ Rejected",
                "Notes": candidate.get("rationale", "")
            })

        df = pd.DataFrame(rows)
        html_table = df.to_html(index=False, classes="ranking-table", border=1)
        
        orig_summary = data.get("Summary", "")
        data["Summary"] = f"<div><h3>Ranked Candidates</h3>{html_table}<p>{orig_summary}</p></div>"

        return jsonify(data)

    except Exception as e:
        return jsonify({"error": "AI Processing/Parsing failed", "details": str(e), "raw_response": raw if 'raw' in locals() else None}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000)
