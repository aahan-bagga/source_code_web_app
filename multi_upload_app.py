from flask import Flask, request, jsonify
from flask_cors import CORS
from google import genai
from google.genai import types
from sentence_transformers import SentenceTransformer, util
from docx import Document
import fitz
import os
import pathlib
import json

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024
CORS(app)

os.environ["TOKENIZERS_PARALLELISM"] = "false"

# Load SBERT model and Gemini API client
model = SentenceTransformer("all-mpnet-base-v2")
genai_client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))

def extract_docx(filepath):
    doc = Document(filepath)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    for section in doc.sections:
        for para in section.header.paragraphs:
            text.append(para.text)
        for para in section.footer.paragraphs:
            text.append(para.text)
    return "\n".join([t for t in text if t.strip() != ""])

def extract_text_and_part(filepath):
    ext = pathlib.Path(filepath).suffix.lower()
    if ext == ".pdf":
        part = types.Part.from_bytes(data=pathlib.Path(filepath).read_bytes(), mime_type="application/pdf")
        doc = fitz.open(filepath)
        text = "\n".join(page.get_text() for page in doc)
    elif ext == ".docx":
        text = extract_docx(filepath)
        part = types.Part(text=text)
    else:
        raise ValueError("Unsupported file type")
    return text, part

@app.route('/score_resumes_ranked', methods=['POST'])
def score_resumes_ranked():
    if 'resumes' not in request.files or 'job' not in request.files:
        return jsonify({"error": "Missing files"}), 400

    job_file = request.files['job']
    job_path = f"/tmp/{job_file.filename}"
    job_file.save(job_path)
    jd_text, jd_part = extract_text_and_part(job_path)

    resume_files = request.files.getlist('resumes')
    results = []

    for resume_file in resume_files:
        resume_path = f"/tmp/{resume_file.filename}"
        resume_file.save(resume_path)
        resume_text, _ = extract_text_and_part(resume_path)

        resume_vec = model.encode(resume_text, convert_to_tensor=True)
        jd_vec = model.encode(jd_text, convert_to_tensor=True)
        sbert_score = util.cos_sim(resume_vec, jd_vec).item()

        results.append({
            "resume_name": resume_file.filename,
            "resume_text": resume_text,
            "sbert_score": round(sbert_score, 3)
        })

    # === PROMPT ===
    prompt = """
You are TalentMatchAI, a hiring expert for tech roles in Indian IT services.
You are given 5 resumes and a single job description. Your task is to return a valid JSON object that contains the ranking of these resumes from most to least suitable with the given job description based on the following:

1. *Fitment Score (1–10)*  
   – Based on skills, experience, and qualifications match.

2. *Selection Decision*  
   – “✅ Selected” or “❌ Rejected”  
   – Add 1–2 sentences explaining the decision.

3. *Skill Gap Analysis*  
   – Count required skills: present vs. missing  
   – Rate present skills: Expert / Proficient / Basic  
   – Format as a table:

   | Skill | Required? | Present? | Depth (Exp/Prof/Basic) |

4. *Relevant Experience Summary*  
   – Total years + domain and tools/tech used.

5. *Skill Presence*  
   – List key JD skills: YES if present, NO if not.

6. *Suggested Domains* (if selected)  
   – 2–3 IT service domains (e.g., BFSI, E‑Commerce) with 1-line reasoning each.


Include the SBERT similarity score for each candidate in your evaluation.
Output in rank order (1 = best fit, 5 = worst fit) in your formatted JSON object and return ONLY THE JSON OBJECT, no other commentary or explanation.
"""

    contents = [types.Part(text=prompt), jd_part]

    for r in results:
        resume_blob = (
            f"\nResume Filename: {r['resume_name']}\n"
            f"SBERT Score: {r['sbert_score']}\n\n"
            f"{r['resume_text']}"
        )
        contents.append(types.Part(text=resume_blob))

    response = genai_client.models.generate_content(
        model="gemini-2.0-flash",
        contents=contents
    )

    try:
        data = json.loads(response.text)
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": "Invalid JSON from Gemini", "raw": response.text}), 500

    #return jsonify({
     #   "Ranking": response.text
    #})

@app.route('/')
def home():
    return 'Resume Fitment Ranking API is live!'

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000)
